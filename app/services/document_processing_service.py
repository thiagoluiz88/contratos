from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from uuid import uuid4

from fastapi import UploadFile
from sqlalchemy.orm import Session

from app.config import MAX_UPLOAD_SIZE_BYTES, UPLOAD_DIR
from app.models import ContractExtraction, ContractFile


ALLOWED_DOCUMENT_EXTENSIONS = {".pdf", ".doc", ".docx", ".png", ".jpg", ".jpeg"}
DOCUMENT_TYPES = {"contrato", "aditivo", "tabela", "anexo", "outro"}
PROCESSING_PENDING_VALIDATION = "aguardando_validacao"
PROCESSING_ERROR = "erro"
PROCESSING_IN_PROGRESS = "em_processamento"
PROCESSING_TEXT_EXTRACTED = "texto_extraido"
TEXT_PREVIEW_LIMIT = 4000


class DocumentProcessingError(ValueError):
    pass


@dataclass(slots=True)
class ProcessedDocument:
    contract_file: ContractFile
    extraction: ContractExtraction


@dataclass(slots=True)
class RawTextExtractionResult:
    text: str | None
    method: str | None
    warnings: list[str]
    page_count: int | None = None
    confidence: float | None = None


def empty_extraction_payload() -> dict:
    return {
        "contrato": {
            "operadora": None,
            "numero_contrato": None,
            "tipo_contrato": None,
            "data_inicio": None,
            "data_fim": None,
            "data_base_reajuste": None,
            "indice_reajuste": None,
            "percentual_reajuste": None,
        },
        "clausulas_criticas": {
            "prazo_faturamento": None,
            "prazo_recurso_glosa": None,
            "regras_glosa": None,
            "regras_autorizacao": None,
            "multas": None,
            "auditoria": None,
        },
        "condicoes_contratuais": [
            {
                "categoria": None,
                "item": None,
                "descricao": None,
                "valor": None,
                "unidade": None,
                "vigencia_inicio": None,
                "vigencia_fim": None,
            }
        ],
    }


def normalize_document_type(value: str | None) -> str:
    normalized = (value or "outro").strip().lower()
    return normalized if normalized in DOCUMENT_TYPES else "outro"


def validate_upload_metadata(file: UploadFile) -> tuple[str, str]:
    original_filename = Path(file.filename or "documento").name
    extension = Path(original_filename).suffix.lower()
    if extension not in ALLOWED_DOCUMENT_EXTENSIONS:
        raise DocumentProcessingError("Formato não suportado. Envie PDF, DOC, DOCX, PNG, JPG ou JPEG.")
    return original_filename, extension


def base_extraction_payload(*, raw_text_available: bool) -> dict:
    payload = empty_extraction_payload()
    payload["raw_text_available"] = raw_text_available
    payload["candidate_sections"] = []
    payload["pending_ai_analysis"] = True
    return payload


def build_text_preview(text: str | None) -> str | None:
    if not text:
        return None
    compact = "\n".join(line.rstrip() for line in text.strip().splitlines())
    return compact[:TEXT_PREVIEW_LIMIT]


def build_processing_note(result: RawTextExtractionResult) -> str:
    if result.text:
        return "Texto bruto extraído automaticamente para apoio à conferência. Dados estruturados dependem de validação humana."
    if result.warnings:
        return " ".join(result.warnings)
    return "Documento preparado para validação humana. Nenhuma IA externa foi executada."


async def process_uploaded_document(
    db: Session,
    *,
    contract_id: int,
    document_type: str,
    file: UploadFile,
    username: str | None,
) -> ProcessedDocument:
    original_filename, extension = validate_upload_metadata(file)

    document_type = normalize_document_type(document_type)
    stored_path, file_size = await _save_upload(file, extension)

    contract_file = ContractFile(
        contract_id=contract_id,
        file_type=document_type,
        document_type=document_type,
        original_filename=original_filename,
        stored_filepath=str(stored_path),
        mime_type=file.content_type,
        file_size_bytes=file_size,
        extraction_status="pendente",
        processing_status=PROCESSING_IN_PROGRESS,
        uploaded_by=username,
    )
    db.add(contract_file)
    db.flush()

    try:
        raw_text_result = extract_raw_text(stored_path, extension)
        extracted_text = raw_text_result.text or None
        text_status = PROCESSING_TEXT_EXTRACTED if extracted_text else PROCESSING_PENDING_VALIDATION
        extraction = ContractExtraction(
            contract_file_id=contract_file.id,
            contract_id=contract_id,
            extraction_status=text_status,
            extracted_json=base_extraction_payload(raw_text_available=bool(extracted_text)),
            extracted_text=extracted_text,
            extracted_text_preview=build_text_preview(extracted_text),
            extraction_method=raw_text_result.method,
            extraction_warnings="\n".join(raw_text_result.warnings) if raw_text_result.warnings else None,
            page_count=raw_text_result.page_count,
            character_count=len(extracted_text or ""),
            confidence_score=raw_text_result.confidence,
            extraction_source=raw_text_result.method or "manual",
            created_by=username,
            review_status="pendente",
        )
        db.add(extraction)
        contract_file.extracted_text = extracted_text
        contract_file.extraction_method = raw_text_result.method
        contract_file.processing_status = PROCESSING_PENDING_VALIDATION
        contract_file.extraction_status = text_status
        contract_file.processed_at = datetime.utcnow()
        contract_file.notes = build_processing_note(raw_text_result)
        db.flush()
        return ProcessedDocument(contract_file=contract_file, extraction=extraction)
    except Exception as exc:
        contract_file.processing_status = PROCESSING_ERROR
        contract_file.extraction_status = PROCESSING_ERROR
        contract_file.error_message = "Falha ao preparar validação do documento."
        db.flush()
        raise DocumentProcessingError("Não foi possível preparar o documento para validação.") from exc


async def _save_upload(file: UploadFile, extension: str) -> tuple[Path, int]:
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    stored_path = UPLOAD_DIR / f"{uuid4().hex}{extension}"
    file_size = 0

    try:
        with stored_path.open("xb") as destination:
            while chunk := await file.read(1024 * 1024):
                file_size += len(chunk)
                if file_size > MAX_UPLOAD_SIZE_BYTES:
                    raise DocumentProcessingError(f"Arquivo excede o limite de {MAX_UPLOAD_SIZE_BYTES // (1024 * 1024)} MB.")
                destination.write(chunk)
        _validate_basic_signature(stored_path, extension)
    except Exception:
        try:
            stored_path.unlink(missing_ok=True)
        except OSError:
            pass
        raise

    return stored_path, file_size


def _validate_basic_signature(stored_path: Path, extension: str) -> None:
    header = stored_path.read_bytes()[:16]
    if extension == ".pdf" and not header.startswith(b"%PDF-"):
        raise DocumentProcessingError("O conteúdo do arquivo não corresponde a um PDF válido.")
    if extension == ".png" and not header.startswith(b"\x89PNG\r\n\x1a\n"):
        raise DocumentProcessingError("O conteúdo do arquivo não corresponde a um PNG válido.")
    if extension in {".jpg", ".jpeg"} and not header.startswith(b"\xff\xd8\xff"):
        raise DocumentProcessingError("O conteúdo do arquivo não corresponde a uma imagem JPEG válida.")
    if extension == ".docx" and not header.startswith(b"PK\x03\x04"):
        raise DocumentProcessingError("O conteúdo do arquivo não corresponde a um DOCX válido.")


def extract_raw_text(stored_path: Path, extension: str) -> RawTextExtractionResult:
    if extension == ".pdf":
        return _extract_pdf_text(stored_path)
    if extension == ".docx":
        return _extract_docx_text(stored_path)
    if extension == ".doc":
        return RawTextExtractionResult(
            text=None,
            method="doc_legacy_unsupported",
            warnings=["Arquivo DOC legado salvo. Conversão para DOCX/PDF é necessária para extração automática nesta etapa."],
        )
    if extension in {".png", ".jpg", ".jpeg"}:
        return _extract_image_text(stored_path)
    return RawTextExtractionResult(text=None, method=None, warnings=["Formato salvo, mas sem extrator de texto configurado."])


def _extract_pdf_text(stored_path: Path) -> RawTextExtractionResult:
    import pdfplumber

    texts: list[str] = []
    page_count = 0
    try:
        with pdfplumber.open(str(stored_path)) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                text = page.extract_text() or ""
                if text.strip():
                    texts.append(text.strip())
    except Exception as exc:
        return RawTextExtractionResult(
            text=None,
            method="pdf_text",
            page_count=page_count or None,
            warnings=[f"Não foi possível ler a camada de texto do PDF: {exc}"],
        )

    extracted = "\n\n".join(texts).strip()
    if extracted:
        return RawTextExtractionResult(
            text=extracted,
            method="pdf_text",
            warnings=[],
            page_count=page_count,
            confidence=0.96,
        )

    ocr_result = _extract_pdf_text_via_ocr(stored_path)
    if ocr_result.text:
        ocr_result.page_count = ocr_result.page_count or page_count
        return ocr_result
    warnings = ["PDF sem texto digital detectável."]
    warnings.extend(ocr_result.warnings)
    return RawTextExtractionResult(text=None, method="pdf_text", warnings=warnings, page_count=page_count)


def _extract_pdf_text_via_ocr(stored_path: Path) -> RawTextExtractionResult:
    try:
        from pdf2image import convert_from_path
        from pdf2image.exceptions import PDFInfoNotInstalledError, PDFPageCountError, PDFSyntaxError
    except ImportError:
        return RawTextExtractionResult(
            text=None,
            method="pdf_ocr_unavailable",
            warnings=["OCR local não configurado. Documento enviado, mas texto não extraído automaticamente."],
        )

    try:
        pages = convert_from_path(str(stored_path), dpi=220)
    except (PDFInfoNotInstalledError, PDFPageCountError, PDFSyntaxError, OSError) as exc:
        return RawTextExtractionResult(
            text=None,
            method="pdf_ocr_unavailable",
            warnings=[f"OCR local não configurado ou Poppler indisponível: {exc}"],
        )

    try:
        result = _extract_images_text_from_objects(pages, method="pdf_ocr")
        result.page_count = len(pages)
        return result
    finally:
        for image in pages:
            try:
                image.close()
            except Exception:
                pass


def _extract_docx_text(stored_path: Path) -> RawTextExtractionResult:
    from docx import Document

    document = Document(str(stored_path))
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    extracted = "\n".join(blocks).strip()
    if not extracted:
        return RawTextExtractionResult(
            text=None,
            method="docx",
            warnings=["DOCX válido, mas sem texto legível detectado."],
        )
    return RawTextExtractionResult(text=extracted, method="docx", warnings=[], confidence=0.98)


def _extract_image_text(stored_path: Path) -> RawTextExtractionResult:
    try:
        from PIL import Image
    except ImportError:
        return RawTextExtractionResult(
            text=None,
            method="image_ocr_unavailable",
            warnings=["OCR local não configurado. Documento enviado, mas texto não extraído automaticamente."],
        )

    try:
        with Image.open(stored_path) as image:
            return _extract_images_text_from_objects([image.copy()], method="image_ocr")
    except Exception as exc:
        return RawTextExtractionResult(
            text=None,
            method="image_ocr",
            warnings=[f"Não foi possível preparar a imagem para OCR: {exc}"],
        )


def _extract_images_text_from_objects(images, method: str) -> RawTextExtractionResult:
    try:
        import pytesseract
    except ImportError:
        return RawTextExtractionResult(
            text=None,
            method=f"{method}_unavailable",
            warnings=["OCR local não configurado. Documento enviado, mas texto não extraído automaticamente."],
        )

    texts: list[str] = []
    confidences: list[float] = []
    for image in images:
        try:
            data = pytesseract.image_to_data(image, lang="por", output_type=pytesseract.Output.DICT)
        except pytesseract.TesseractNotFoundError:
            return RawTextExtractionResult(
                text=None,
                method=f"{method}_unavailable",
                warnings=["OCR local não configurado. Documento enviado, mas texto não extraído automaticamente."],
            )
        except pytesseract.TesseractError as exc:
            return RawTextExtractionResult(text=None, method=method, warnings=[f"OCR local não conseguiu ler o arquivo: {exc}"])
        finally:
            try:
                image.close()
            except Exception:
                pass

        words = []
        for index, word in enumerate(data.get("text", [])):
            word = (word or "").strip()
            if not word:
                continue
            words.append(word)
            try:
                confidence = float(data["conf"][index])
                if confidence >= 0:
                    confidences.append(confidence)
            except Exception:
                pass
        if words:
            texts.append(" ".join(words))

    extracted = "\n\n".join(texts).strip()
    if not extracted:
        return RawTextExtractionResult(
            text=None,
            method=method,
            warnings=["OCR local executado, mas nenhum texto legível foi detectado."],
        )
    avg_confidence = (sum(confidences) / len(confidences) / 100.0) if confidences else 0.70
    return RawTextExtractionResult(text=extracted, method=method, warnings=[], confidence=round(avg_confidence, 2))


def apply_approved_extraction_to_contract(*args, **kwargs) -> None:
    """Hook futuro para aplicar dados aprovados ao cadastro do contrato."""
    return None


def apply_extracted_terms_to_contract_terms(*args, **kwargs) -> None:
    """Hook futuro para versionar condições aprovadas em contract_terms."""
    return None
